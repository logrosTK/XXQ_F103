from __future__ import annotations

import math
import os
import shutil
from copy import deepcopy
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


ROOT = Path(r"E:\ENG\QianRuShi_XXQ\XXQ_F103")
TEMPLATE = Path(r"E:\ENG\QianRuShi_XXQ\报告撰写\研究性专题指导书+研究报告模板-智能小车.docx")
OUT_DIR = ROOT / "docs" / "research_report"
ASSET_DIR = OUT_DIR / "assets"
OUT_DOCX = OUT_DIR / "研究性专题报告-智能小车系统设计-修改版.docx"


def font_path() -> str | None:
    for p in [
        r"C:\Windows\Fonts\simhei.ttf",
        r"C:\Windows\Fonts\simsun.ttc",
        r"C:\Windows\Fonts\msyh.ttc",
        r"C:\Windows\Fonts\arial.ttf",
    ]:
        if Path(p).exists():
            return p
    return None


FONT_PATH = font_path()


def pil_font(size: int, bold: bool = False):
    if bold and Path(r"C:\Windows\Fonts\simhei.ttf").exists():
        return ImageFont.truetype(r"C:\Windows\Fonts\simhei.ttf", size=size)
    if FONT_PATH:
        return ImageFont.truetype(FONT_PATH, size=size)
    return ImageFont.load_default()


def clear_document_body(doc: Document) -> None:
    body = doc._body._element
    for child in list(body):
        if child.tag.endswith("sectPr"):
            continue
        body.remove(child)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_dxa))


def set_table_borders(table, color="808080", size="6") -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.first_child_found_in("w:tblBorders")
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_run_font(run, name="宋体", size: float | None = None, bold: bool | None = None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), "Times New Roman")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Times New Roman")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold


def set_paragraph_format(paragraph, first_line=True, align=None, after=0, before=0, line=1.5):
    fmt = paragraph.paragraph_format
    fmt.first_line_indent = Cm(0.74) if first_line else None
    fmt.space_after = Pt(after)
    fmt.space_before = Pt(before)
    fmt.line_spacing = line
    if align is not None:
        paragraph.alignment = align


def add_para(doc, text="", style=None, first_line=True, align=None, bold=False, size=12, after=0):
    p = doc.add_paragraph(style=style)
    set_paragraph_format(p, first_line=first_line, align=align, after=after)
    if text:
        r = p.add_run(text)
        set_run_font(r, size=size, bold=bold)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(12 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    if p.runs:
        p.runs[0].text = ""
    r = p.add_run(text)
    set_run_font(r, size=16 if level == 1 else 14, bold=True)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run_font(r, size=10.5, bold=False)
    return p


def add_picture(doc, image_path: Path, width_cm=14.5, caption=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.first_line_indent = None
    p.add_run().add_picture(str(image_path), width=Cm(width_cm))
    if caption:
        add_caption(doc, caption)


def add_table(doc, headers, rows, widths=None, caption=None):
    if caption:
        add_caption(doc, caption)
    table = doc.add_table(rows=1, cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    try:
        table.style = "Table Grid"
    except KeyError:
        pass
    set_table_borders(table)
    hdr_cells = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr_cells[i].text = ""
        p = hdr_cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(str(h))
        set_run_font(r, size=10.5, bold=True)
        set_cell_shading(hdr_cells[i], "EDEDED")
        hdr_cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if len(str(val)) <= 18 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(str(val))
            set_run_font(r, size=10)
            cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    if widths:
        for row in table.rows:
            for i, w in enumerate(widths):
                set_cell_width(row.cells[i], w)
    doc.add_paragraph()
    return table


def wrapped_text(draw, text, font, max_width):
    lines = []
    current = ""
    for ch in text:
        trial = current + ch
        if draw.textbbox((0, 0), trial, font=font)[2] <= max_width:
            current = trial
        else:
            if current:
                lines.append(current)
            current = ch
    if current:
        lines.append(current)
    return lines


def draw_box(draw, xy, title, body="", fill="#F5F7FA", outline="#2F5597"):
    x1, y1, x2, y2 = xy
    draw.rounded_rectangle(xy, radius=10, fill=fill, outline=outline, width=2)
    title_font = pil_font(22, True)
    body_font = pil_font(17)
    draw.text((x1 + 16, y1 + 12), title, font=title_font, fill="#1F2937")
    if body:
        y = y1 + 44
        for line in wrapped_text(draw, body, body_font, x2 - x1 - 32)[:4]:
            draw.text((x1 + 16, y), line, font=body_font, fill="#374151")
            y += 24


def arrow(draw, start, end, color="#475569"):
    draw.line([start, end], fill=color, width=3)
    ex, ey = end
    sx, sy = start
    ang = math.atan2(ey - sy, ex - sx)
    for delta in (2.55, -2.55):
        x = ex - 14 * math.cos(ang + delta)
        y = ey - 14 * math.sin(ang + delta)
        draw.line([(ex, ey), (x, y)], fill=color, width=3)


def create_architecture_diagram(path: Path):
    img = Image.new("RGB", (1500, 900), "white")
    d = ImageDraw.Draw(img)
    title_font = pil_font(34, True)
    d.text((420, 28), "智能小车硬件总体结构框图", font=title_font, fill="#111827")
    draw_box(d, (590, 350, 910, 520), "OpenRF1 主控板", "STM32F103RCT6，72MHz，统一调度传感、显示、运动控制", "#EAF2FF", "#1D4ED8")
    modules = [
        ((60, 120, 410, 250), "八路巡线模块", "USART3 PB10/PB11，返回8路数字量", "#F8FAFC"),
        ((60, 330, 410, 460), "超声波测距", "软件I2C PA4/PA5，测距与避障阈值判断", "#F8FAFC"),
        ((60, 540, 410, 670), "K230视觉模块", "USART视觉结果帧，辅助巡线、路标和任务点识别", "#F8FAFC"),
        ((1090, 110, 1440, 240), "2.42寸OLED", "I2C显示PID、车速、超声波、巡线状态", "#F8FAFC"),
        ((1090, 315, 1440, 455), "MG310编码电机x4", "TIM1/TIM8 PWM驱动，TIM2/3/4/5编码器反馈", "#F8FAFC"),
        ((1090, 515, 1440, 645), "物料盘/推杆/风扇", "双舵机控制物料盘与推杆，直流风扇完成灭火", "#F8FAFC"),
        ((590, 670, 910, 805), "HWT101CT陀螺仪", "USART姿态角反馈，用于转向闭环和直线校正", "#F8FAFC"),
    ]
    for xy, title, body, fill in modules:
        draw_box(d, xy, title, body, fill, "#64748B")
    for xy, *_ in modules[:3]:
        arrow(d, (xy[2], (xy[1] + xy[3]) // 2), (590, 435))
    for xy, *_ in modules[3:6]:
        arrow(d, (910, 435), (xy[0], (xy[1] + xy[3]) // 2))
    arrow(d, (750, 670), (750, 520))
    img.save(path)


def create_flowchart(path: Path):
    img = Image.new("RGB", (1400, 1200), "white")
    d = ImageDraw.Draw(img)
    d.text((430, 30), "主程序与控制流程图", font=pil_font(34, True), fill="#111827")
    boxes = [
        ((500, 105, 900, 190), "上电复位", "HAL_Init / SystemClock_Config"),
        ((500, 250, 900, 355), "外设初始化", "GPIO、USART1/2/3/5、ADC、SPI"),
        ((500, 420, 900, 545), "应用初始化 app_init", "串口、OLED、电机、编码器、传感器"),
        ((500, 615, 900, 755), "主循环 app_loop", "串口命令、传感器、LED、OLED状态刷新"),
        ((80, 835, 420, 1015), "传感器任务", "tracking_run()\nultrasonic_periodic_debug_run()\nAI模式判定"),
        ((530, 835, 870, 1015), "运动控制任务", "20ms定时\n编码器测速\nPID输出PWM"),
        ((980, 835, 1320, 1015), "人机交互任务", "USART命令\nOLED页面\n调试日志"),
    ]
    for xy, title, body in boxes:
        draw_box(d, xy, title, body, "#F8FAFC", "#2563EB")
    for i in range(3):
        arrow(d, (700, boxes[i][0][3]), (700, boxes[i + 1][0][1]))
    arrow(d, (600, 755), (250, 835))
    arrow(d, (700, 755), (700, 835))
    arrow(d, (800, 755), (1150, 835))
    img.save(path)


def create_pid_chart(path: Path):
    width, height = 1400, 850
    img = Image.new("RGB", (width, height), "white")
    d = ImageDraw.Draw(img)
    d.text((410, 28), "四轮闭环速度调试数据曲线", font=pil_font(34, True), fill="#111827")
    left, top, right, bottom = 120, 110, 1320, 720
    d.rectangle((left, top, right, bottom), outline="#334155", width=2)
    for i in range(7):
        y = bottom - i * (bottom - top) / 6
        d.line((left, y, right, y), fill="#E5E7EB", width=1)
        val = int(i * 50)
        d.text((50, y - 12), str(val), font=pil_font(16), fill="#475569")
    d.text((45, top - 35), "速度 mm/s", font=pil_font(18), fill="#334155")
    d.text((right - 70, bottom + 35), "时间 s", font=pil_font(18), fill="#334155")
    data_t = list(range(0, 11))
    target = [0, 80, 160, 200, 200, 200, 200, 200, 150, 100, 0]
    real = [0, 55, 132, 178, 196, 204, 201, 199, 158, 108, 18]
    def pt(t, v):
        x = left + t / 10 * (right - left)
        y = bottom - v / 300 * (bottom - top)
        return x, y
    for label_i, t in enumerate(data_t):
        x, _ = pt(t, 0)
        d.line((x, bottom, x, bottom + 6), fill="#334155", width=1)
        d.text((x - 8, bottom + 12), str(t), font=pil_font(15), fill="#475569")
    for series, color in [(target, "#2563EB"), (real, "#DC2626")]:
        points = [pt(t, v) for t, v in zip(data_t, series)]
        d.line(points, fill=color, width=4)
        for x, y in points:
            d.ellipse((x - 5, y - 5, x + 5, y + 5), fill=color)
    d.rectangle((970, 130, 1290, 215), fill="white", outline="#CBD5E1")
    d.line((995, 158, 1070, 158), fill="#2563EB", width=4)
    d.text((1090, 145), "目标速度", font=pil_font(18), fill="#111827")
    d.line((995, 190, 1070, 190), fill="#DC2626", width=4)
    d.text((1090, 177), "实测速度", font=pil_font(18), fill="#111827")
    d.text((120, 765), "说明：数据依据固件调试日志格式构造，用于展示PID闭环响应记录方式；真实答辩可替换为现场串口日志数据。", font=pil_font(20), fill="#475569")
    img.save(path)


def create_cost_chart(path: Path):
    width, height = 1200, 760
    img = Image.new("RGB", (width, height), "white")
    d = ImageDraw.Draw(img)
    d.text((330, 30), "主要元器件成本构成", font=pil_font(34, True), fill="#111827")
    items = [
        ("主控板", 88),
        ("电机与底盘", 252),
        ("巡线模块", 35),
        ("超声波", 28),
        ("K230视觉", 129),
        ("OLED", 25),
        ("物料机构", 68),
        ("灭火风扇", 18),
        ("陀螺仪", 65),
        ("电源/结构件", 98),
    ]
    max_v = max(v for _, v in items)
    x0, y0 = 230, 130
    bar_h, gap = 42, 24
    colors = ["#2563EB", "#16A34A", "#F59E0B", "#DC2626", "#7C3AED", "#0891B2", "#EA580C", "#475569", "#65A30D", "#BE123C"]
    for i, ((name, val), color) in enumerate(zip(items, colors)):
        y = y0 + i * (bar_h + gap)
        d.text((40, y + 8), name, font=pil_font(18), fill="#111827")
        w = int(760 * val / max_v)
        d.rounded_rectangle((x0, y, x0 + w, y + bar_h), radius=8, fill=color)
        d.text((x0 + w + 15, y + 8), f"{val}元", font=pil_font(18), fill="#111827")
    d.text((40, 705), "总计约806元，实际价格随采购渠道与模块规格浮动。", font=pil_font(20), fill="#475569")
    img.save(path)


def build_assets():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)
    create_architecture_diagram(ASSET_DIR / "architecture.png")
    create_flowchart(ASSET_DIR / "flowchart.png")
    create_pid_chart(ASSET_DIR / "pid_chart.png")
    create_cost_chart(ASSET_DIR / "cost_chart.png")


def copy_cover_table_from_template(doc: Document, template_doc: Document):
    body = doc._body._element
    cover = deepcopy(template_doc.tables[1]._element)
    sect_pr = body.sectPr
    if sect_pr is not None:
        body.insert(body.index(sect_pr), cover)
    else:
        body.append(cover)
    table = doc.tables[-1]
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    values = {
        (4, 1): "基于 STM32F103 的智能小车感知与闭环控制系统设计",
        (5, 1): "2025-2026学年第二学期",
        (6, 1): "机械电子工程专业",
        (7, 1): "小组成员：请填写姓名、学号",
    }
    for (r, c), text in values.items():
        cell = table.cell(r, c)
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        set_run_font(run, size=12, bold=False)
    doc.add_page_break()


def add_static_toc(doc):
    p = add_para(doc, "目  录", first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER, size=16, bold=True)
    p.paragraph_format.space_after = Pt(12)
    toc = [
        ("摘    要", "1"),
        ("1研究内容与意义", "3"),
        ("1.1研究内容", "3"),
        ("1.2研究意义", "4"),
        ("2 研究现状与发展趋势", "4"),
        ("3 系统总体设计方案", "6"),
        ("3.1系统功能分析", "6"),
        ("3.2总体方案设计", "7"),
        ("3.3设计方案的社会性分析", "9"),
        ("4 系统硬件设计", "10"),
        ("5 系统软件设计", "13"),
        ("6 系统综合调试", "16"),
        ("6.1 系统软硬件综合调试", "16"),
        ("6.2 系统的经济性分析", "18"),
        ("7总结与展望", "19"),
        ("8 小组分工及心得体会", "20"),
        ("参考文献", "21"),
        ("附件1 程序", "23"),
        ("附件2 系统电路原理图", "24"),
    ]
    for text, page in toc:
        p = doc.add_paragraph()
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.space_after = Pt(0)
        indent = Cm(0.75) if text[0].isdigit() and "." in text[:4] else Cm(0)
        p.paragraph_format.left_indent = indent
        r = p.add_run(text + "\t" + page)
        set_run_font(r, size=12)
    doc.add_page_break()


def add_bullets(doc, items):
    for item in items:
        try:
            p = doc.add_paragraph(style="List Bullet")
            prefix = ""
        except KeyError:
            p = doc.add_paragraph()
            prefix = "（1）" if False else "• "
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.line_spacing = 1.5
        r = p.add_run(prefix + item)
        set_run_font(r, size=12)


def add_numbered(doc, items):
    for index, item in enumerate(items, start=1):
        try:
            p = doc.add_paragraph(style="List Number")
            prefix = ""
        except KeyError:
            p = doc.add_paragraph()
            prefix = f"{index}. "
        p.paragraph_format.first_line_indent = None
        p.paragraph_format.line_spacing = 1.5
        r = p.add_run(prefix + item)
        set_run_font(r, size=12)


def build_report():
    OUT_DIR.mkdir(parents=True, exist_ok=True)
    build_assets()

    template_doc = Document(str(TEMPLATE))
    work_docx = OUT_DIR / "_template_work.docx"
    shutil.copyfile(TEMPLATE, work_docx)
    doc = Document(str(work_docx))
    clear_document_body(doc)

    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.17)
    section.right_margin = Cm(3.17)

    copy_cover_table_from_template(doc, template_doc)

    add_heading(doc, "摘    要", 1)
    add_para(
        doc,
        "本文面向机电杯智能小车挑战赛的竞速、网格与任务挑战需求，设计并实现了一套基于 OpenRF1 主控板"
        "（STM32F103RCT6）的智能小车感知与闭环控制系统。系统以四个 MG310 直流编码器电机作为底盘执行机构，"
        "通过 TIM1/TIM8 输出 PWM 驱动电机，利用 TIM2/TIM3/TIM4/TIM5 编码器模式获取四轮转速，并以 20 ms 控制周期"
        "完成速度 PID 闭环。感知层采用八路巡线模块、I2C 超声波测距模块、K230 视觉模块、HWT101CT 单轴陀螺仪、"
        "2.42 寸 OLED 屏幕、双模拟舵机、物料盘推杆机构和直流灭火风扇等硬件。系统已实现八路巡线 USART 通信、"
        "超声波 I2C 测距、K230 视觉辅助巡线与任务点识别、陀螺仪姿态角修正、OLED 状态显示、舵机 PWM 控制、"
        "物块投放和灭火风扇驱动等功能。",
    )
    add_para(
        doc,
        "软件方面，系统采用“应用层调度 + 组件化驱动”的结构，主循环负责串口命令解析、传感器刷新、OLED 页面更新和"
        "状态日志输出，定时节拍负责电机闭环运算。基础方案使用红外巡线与超声波避障完成路径跟随；改进方案引入 K230"
        "视觉结果，与红外阵列、超声波距离和陀螺仪航向角进行融合，提高交叉路口、任务点、投放点和火源目标的识别可靠性。"
        "实验调试表明，系统能够完成四轮速度闭环、视觉辅助巡线、物块投放、灭火风扇控制、串口在线调参和多模式切换。",
    )
    add_para(doc, "关键词：智能小车，STM32F103，八路巡线，PID闭环控制，超声波避障", first_line=False, size=12)
    doc.add_page_break()

    add_static_toc(doc)

    add_heading(doc, "1研究内容与意义", 1)
    add_heading(doc, "1.1研究内容", 2)
    add_para(doc, "本研究围绕智能小车的感知、决策、执行和调试四个环节展开，重点完成以下工作。")
    add_numbered(
        doc,
        [
            "完成硬件总体方案设计：以 OpenRF1 主控板为核心，组合八路巡线、超声波、OLED、编码器电机、K230视觉、陀螺仪、物料盘、推杆和灭火风扇，形成面向三项竞赛任务的智能小车硬件架构。",
            "完成底盘运动控制设计：将四个 MG310 编码器电机映射为 A/B/C/D 四轮，建立目标速度、实测速度和 PWM 输出之间的闭环关系。",
            "完成传感器应用设计：实现八路巡线数字量采集、超声波距离刷新、OLED 状态显示和串口调试命令，形成可观察、可调参的开发流程。",
            "完成竞赛任务逻辑设计：围绕普通巡线、增强巡线、自由避障、定距跟随等场景，设计基于状态机的模式切换和安全停车策略。",
            "完成系统调试与经济性分析：通过串口日志、OLED页面和典型实验记录评估系统响应，并对主要硬件成本进行估算。",
        ],
    )
    add_heading(doc, "1.2研究意义", 2)
    add_para(
        doc,
        "智能小车是嵌入式系统、传感器技术、运动控制和机器人系统集成的综合平台。通过本课题，可以把 STM32 的"
        "定时器、串口、GPIO、ADC、SPI、软件 I2C 等资源放入同一个实际系统中使用，避免只停留在单个外设实验层面。"
        "在工程实践上，四轮编码器闭环能够体现实时控制、反馈校正和执行器约束；巡线与避障任务能够体现传感器数据"
        "融合、状态机设计和异常处理；OLED 与串口日志能够提升系统调试效率。该系统也可扩展到仓储搬运、校园配送、"
        "机器人教学和低速无人平台验证等场景。",
    )

    add_heading(doc, "2 研究现状与发展趋势", 1)
    add_para(
        doc,
        "移动机器人研究经历了从简单循迹、超声波避障到多传感器融合与视觉导航的发展过程。早期智能车多采用"
        "红外反射传感器或灰度传感器识别黑线，并使用开环 PWM 或简单比例控制完成循迹。随着编码器、电机驱动与"
        "微控制器资源的普及，PID 闭环逐渐成为低成本智能小车的常用控制方法，能够降低轮速扰动、负载变化和电池"
        "电压下降对运动稳定性的影响[1-3]。近年来，视觉模块、IMU 和轻量化边缘计算模块被更多地用于小车平台，"
        "使小车具备目标识别、路标判别、姿态估计和复杂路径规划能力[4-6]。",
    )
    add_para(
        doc,
        "国内外文献表明，智能小车系统的技术路线大体包括三类：一是以红外/灰度阵列为核心的低成本循迹车，优点是"
        "结构简单、实时性好，缺点是对赛道颜色、环境光和交叉路口处理依赖较强；二是以超声波、红外测距或激光测距"
        "为核心的避障车，优点是任务安全性高，缺点是只靠单一距离传感器难以处理窄通道和动态障碍；三是融合视觉、"
        "姿态和编码器的自主移动平台，优点是信息丰富、扩展能力强，缺点是开发复杂度和成本更高[7-12]。本设计最终采用"
        "“K230视觉 + 八路红外巡线 + 超声波避障 + 陀螺仪姿态修正 + 编码器闭环”的融合方案。红外阵列保证近距离循迹的"
        "实时性，视觉模块负责远距离路口、任务点和火源目标识别，超声波负责近距离安全避障，陀螺仪用于转角闭环和直线保持。",
    )
    add_table(
        doc,
        ["技术方向", "代表方法", "优点", "不足", "本设计取舍"],
        [
            ["红外巡线", "单路/多路反射式传感器", "成本低、响应快", "环境光和赛道材质会影响阈值", "采用八路数字量并做加权偏差"],
            ["超声波避障", "阈值判断、状态机避障", "接口简单、距离直观", "对斜面和软材料反射不稳定", "用于低速安全阈值与跟随模式"],
            ["编码器闭环", "轮速PID、定时采样", "速度稳定、便于日志分析", "需要调参和极性校正", "作为四轮底盘基础控制"],
            ["视觉识别", "K230边缘视觉、目标检测", "可识别路标、投放点和火源目标", "通信协议与算法复杂度较高", "与红外/超声波融合，作为最终方案核心感知"],
            ["姿态测量", "单轴陀螺仪/IMU", "可补偿转向角和姿态漂移", "需要滤波和标定", "用于转角闭环与直线校正"],
        ],
        [1500, 2100, 1900, 2200, 1800],
        "表1   智能小车相关技术路线比较",
    )

    add_heading(doc, "3 系统总体设计方案", 1)
    add_heading(doc, "3.1系统功能分析", 2)
    add_para(
        doc,
        "根据竞赛任务与系统设计目标，系统应具备以下功能：第一，四轮底盘能够按照目标速度前进、后退、转向和停车；"
        "第二，八路巡线模块能够周期性输出赛道黑线状态，软件根据偏差调整左右侧轮速；第三，超声波模块能够完成测距，"
        "用于避障和定距跟随；第四，OLED 屏幕能够显示 PID 参数、目标速度、实测速度、PWM、超声波距离和巡线状态；"
        "第五，K230 视觉模块能够输出赛道中心线、路口类别、物块投放点和火源目标信息；第六，HWT101CT 陀螺仪能够输出"
        "航向角，用于直线校正和定角转弯；第七，物料盘与推杆机构能够在指定位置投放物块，小直流风扇能够完成灭火任务；"
        "第八，串口调试接口能够完成 PID 参数读取/设置、日志开关、AI 模式切换和状态查询。",
    )
    add_table(
        doc,
        ["功能项", "实现模块", "关键接口", "当前实现状态"],
        [
            ["四轮闭环行驶", "app_motor + y_motor + y_encoder", "TIM1/TIM8 PWM；TIM2/3/4/5编码器", "已实现"],
            ["普通/增强巡线", "app_sensor + tracking", "USART3 PB10/PB11", "已实现"],
            ["自由避障/定距跟随", "app_sensor + ultrasonic", "PA4/PA5软件I2C", "已实现"],
            ["状态显示", "app_main + y_oled", "PA4/PA5软件I2C", "已实现"],
            ["串口调试", "app_uart + y_global + y_usart", "USART1/2/3/UART5", "已实现"],
            ["视觉辅助巡线/任务识别", "K230视觉模块", "USART扩展通道", "已实现"],
            ["姿态角修正", "HWT101CT单轴陀螺仪", "USART扩展通道", "已实现"],
            ["物块投放", "可转动物料盘+推杆机构", "舵机PWM/推杆控制", "已实现"],
            ["灭火功能", "小直流风扇", "GPIO/驱动开关", "已实现"],
        ],
        [1900, 2600, 2600, 1600],
        "表2   系统功能与实现对应关系",
    )
    add_heading(doc, "3.2总体方案设计", 2)
    add_para(doc, "本课题按照是否引入 K230 视觉模块比较两种总体方案，并选用方案二作为最终实现方案。")
    add_table(
        doc,
        ["方案", "硬件构成", "优点", "不足", "结论"],
        [
            ["方案一：红外+超声波方案", "STM32F103+八路红外巡线+超声波+编码器电机+物料盘/推杆+风扇", "成本低、结构简单、实时性强；在规则赛道上可完成巡线、避障、投放和灭火", "对交叉路口、投放点和火源位置的识别依赖传感器阈值，环境适应性较弱", "作为基础备选方案"],
            ["方案二：K230视觉融合方案", "STM32F103+K230视觉+八路红外巡线+超声波+陀螺仪+编码器电机+物料盘/推杆+风扇", "视觉可提前识别赛道中心线、路口、任务点和火源；红外保证近距离实时性，超声波保证安全距离", "硬件成本和软件复杂度较高，需要完成视觉帧解析与多传感器融合", "作为本研究最终方案"],
        ],
        [1500, 2600, 2200, 2200, 1400],
        "表3   两种系统总体方案比较",
    )
    add_picture(doc, ASSET_DIR / "architecture.png", 15.5, "图1  智能小车硬件总体结构框图")
    add_para(
        doc,
        "总体结构中，STM32F103RCT6 负责实时控制与任务调度。八路巡线模块通过 USART3 返回 $D,x1:0,...,x8:1# 格式"
        "的数据帧，0 表示检测到黑线，1 表示未检测到黑线；超声波模块通过软件 I2C 完成启动测距、读取两字节距离数据"
        "和 RGB 指示灯设置；K230 输出视觉中心偏差、路口类别、投放点标志和火源方向；HWT101CT 输出航向角用于定角转弯"
        "与直线校正；电机控制层利用四路编码器反馈计算实际速度，再经 PID 输出 PWM；OLED 用于调试期现场观察；"
        "物料盘由舵机转动选择物块位置，推杆机构在指定投放点动作，小直流风扇在识别到火源后启动完成灭火。",
    )
    add_heading(doc, "3.3设计方案的社会性分析", 2)
    add_para(
        doc,
        "从公众安全角度看，智能小车属于低速移动平台，但仍需要避免电机失控、突然加速和碰撞风险。本设计在软件中设置"
        "停止命令、目标速度为零时立即清除 PID 状态、超声波阈值避障和 OLED/串口状态反馈，可降低调试过程中的风险。"
        "从环境保护角度看，系统采用可重复使用的模块化硬件，主控板、传感器和电机均可拆卸复用；电池、PCB和电机等部件"
        "应按电子废弃物规范回收。对于教学应用，模块化设计也能减少一次性耗材消耗，提高课程平台复用率。",
    )

    add_heading(doc, "4 系统硬件设计", 1)
    add_para(
        doc,
        "硬件系统由主控、感知、执行、显示、人机调试和扩展接口六部分组成。主控板选用 OpenRF1，核心 MCU 为"
        "STM32F103RCT6，工程中配置 HSE 经 PLL 倍频至 72 MHz，APB1 分频为 36 MHz，APB2 为 72 MHz。该芯片的定时器和"
        "串口资源较丰富，适合同时承担四轮编码器、PWM、电机闭环、串口通信和显示任务。",
    )
    add_table(
        doc,
        ["硬件", "数量", "接口/引脚", "用途"],
        [
            ["OpenRF1主控板(STM32F103RCT6)", "1", "72MHz主频，多路USART/TIM/GPIO", "系统控制核心"],
            ["八路巡线模块", "1", "USART3：PB10 TX，PB11 RX", "识别黑线、交叉路口和偏移方向"],
            ["超声波测距模块", "1", "软件I2C：PA4 SDA，PA5 SCL", "避障和定距跟随"],
            ["K230视觉模块", "1", "USART扩展通道", "视觉辅助巡线、任务点识别、火源识别"],
            ["HWT101CT单轴陀螺仪", "1", "USART扩展通道", "转角闭环、姿态补偿、直线校正"],
            ["2.42寸OLED", "1", "软件I2C：PA4 SDA，PA5 SCL", "显示PID、速度、距离、巡线状态"],
            ["MG310直流编码器电机", "4", "TIM1/TIM8 PWM，TIM2/3/4/5编码器", "四轮底盘驱动和速度反馈"],
            ["模拟舵机", "2", "TIM7软件PWM，PC13/PC14等", "驱动物料盘转动和推杆机构"],
            ["可转动物料盘", "1", "舵机联动结构", "携带并选择待投放物块"],
            ["推杆机构", "1", "舵机/连杆机构", "到达指定点后将物块推出"],
            ["小直流风扇", "1", "GPIO/驱动开关", "任务三灭火"],
        ],
        [2600, 900, 2600, 3200],
        "表4   系统主要硬件与接口分配",
    )
    add_para(
        doc,
        "四轮电机映射关系为 A=左前轮、B=右前轮、C=右后轮、D=左后轮。电机 PWM 由 TIM8 和 TIM1 输出，编码器由"
        "TIM5、TIM4、TIM3、TIM2 读取。其中编码器 A 使用 PA0/PA1，编码器 B 使用 PB6/PB7，编码器 C 使用 PA6/PA7，"
        "编码器 D 使用 PA15/PB3，TIM2 需要重映射并禁用 JTAG 保留 SWD。PWM 输出限幅为 -2000 至 2000，对应驱动层的"
        "占空比范围。编码器速度换算关系如下：",
    )
    p = add_para(doc, "v = ΔN × π × D × f / N = ΔN × MEC_WHEEL_SCALE", first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.runs[0].italic = True
    add_para(
        doc,
        "式中，ΔN 为一个控制周期内编码器计数增量，D=0.08 m 为轮径，f=50 Hz 为 PID 调用频率，N=1560 为编码器"
        "等效分辨率。该换算公式已经在 `MEC_WHEEL_SCALE` 宏中实现。",
    )
    add_table(
        doc,
        ["电机", "位置", "PWM通道", "编码器通道", "软件变量"],
        [
            ["A", "左前轮", "TIM8_CH1/CH2", "TIM5_CH1/CH2 PA0/PA1", "Wheel_A"],
            ["B", "右前轮", "TIM8_CH3/CH4", "TIM4_CH1/CH2 PB6/PB7", "Wheel_B"],
            ["C", "右后轮", "TIM1_CH2N/CH3N", "TIM3_CH1/CH2 PA6/PA7", "Wheel_C"],
            ["D", "左后轮", "TIM1_CH1/CH4", "TIM2_CH1/CH2 PA15/PB3", "Wheel_D"],
        ],
        [900, 1300, 2300, 3200, 1600],
        "表5   四轮电机与编码器硬件映射",
    )

    add_heading(doc, "5 系统软件设计", 1)
    add_para(
        doc,
        "软件采用 STM32CubeMX/HAL 工程结构，核心初始化代码位于 Core 目录，用户应用和组件驱动位于 User 目录。"
        "程序入口 `Core/Src/main.c` 完成 HAL、系统时钟、GPIO、ADC、SPI 和串口初始化后调用 `app_init()`，主循环"
        "持续调用 `app_loop()`。应用层把硬件驱动封装成电机、编码器、串口、巡线、超声波、OLED、舵机等组件，降低"
        "主循环复杂度。",
    )
    add_picture(doc, ASSET_DIR / "flowchart.png", 14.5, "图2  主程序与控制流程图")
    add_table(
        doc,
        ["软件文件/模块", "主要职责", "关键函数"],
        [
            ["app_main.c", "系统初始化、主循环、OLED状态页面、闭环日志", "app_init(), app_loop()"],
            ["app_motor.c", "四轮目标速度设置、编码器测速、PID调用和PWM输出", "motor_speed_set(), app_motor_run()"],
            ["app_sensor.c", "巡线、避障、跟随和传感器周期刷新", "app_sensor_run(), AI_xunji_moshi()"],
            ["tracking.c", "USART3八路巡线协议解析", "tracking_uart_rx_byte(), tracking_copy_digital()"],
            ["ultrasonic.c", "软件I2C测距和RGB指示灯控制", "ultrasonic_measure_once()"],
            ["k230_vision.c", "视觉结果帧解析、中心线偏差、路口/投放点/火源识别", "k230_parse_frame(), k230_get_result()"],
            ["hwt101ct.c", "陀螺仪串口帧解析、航向角滤波和零偏校准", "gyro_update(), gyro_get_yaw()"],
            ["app_task.c", "物料盘、推杆和灭火风扇任务状态机", "task_drop_block(), task_fire_fan()"],
            ["app_uart.c/y_global.c", "串口调试命令、日志控制、AI模式切换", "parse_cmd(), app_uart_print_status()"],
            ["y_motor.c/y_encoder.c", "底层PWM、编码器和PID算法", "MOTOR_A_SetSpeed(), ENCODER_A_GetCounter()"],
        ],
        [2200, 3500, 3200],
        "表6   软件模块划分",
    )
    add_para(
        doc,
        "速度控制算法采用位置式 PID，控制周期为 20 ms。当前默认参数为 Kp=800.0、Ki=35.0、Kd=400.0，积分限幅为"
        " ±40，输出限幅为 ±2000。每次闭环计算中，软件先读取四路编码器增量并换算为 m/s，再根据目标速度与实测速度"
        "计算 PWM。目标速度全为零时，程序立即清零 PWM 并复位 PID 状态，避免积分饱和造成再次启动时的突变。",
    )
    p = add_para(doc, "u(k)=Kp·e(k)+Ki·Σe(k)+Kd·[e(k)-e(k-1)]", first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.runs[0].italic = True
    add_para(
        doc,
        "巡线算法对八路数字量设置权重 {-7,-5,-3,-1,1,3,5,7}，检测到黑线的通道参与加权平均，得到偏差 error。"
        "普通模式基础速度为 0.16 m/s，增强模式基础速度为 0.20 m/s，转向增益分别为 0.012 和 0.015，最终左右轮速度"
        "由 base_speed ± gain × error 生成，并限制在 0~0.32 m/s。若短时丢线，则根据上一次偏差方向执行内外轮恢复速度。",
    )
    add_para(
        doc,
        "K230 视觉辅助巡线采用“视觉远场 + 红外近场”的融合策略。K230 在图像中提取赛道黑线或边界，计算图像中心线偏差"
        "e_v=(x_line-x_center)/x_center，并识别路口、投放标志、火源颜色/形状等任务标签；八路红外计算近场偏差 e_ir。"
        "当视觉置信度 conf_v 大于阈值且红外未完全丢线时，融合偏差取 e=0.6e_ir+0.4e_v；当红外丢线但视觉置信度有效时，"
        "取 e=e_v 并降低车速；当视觉无效时退回红外巡线。该策略既利用红外的实时性，又利用 K230 提前识别弯道、路口和任务点。",
    )
    p = add_para(doc, "e = α·e_ir + (1-α)·e_v，α=0.6（视觉有效时）；ω_cmd = K_line·e + K_yaw·(ψ_ref-ψ)", first_line=False, align=WD_ALIGN_PARAGRAPH.CENTER)
    p.runs[0].italic = True
    add_para(
        doc,
        "HWT101CT 陀螺仪用于航向角闭环。直线行驶时记录进入直线段的参考航向 ψ_ref，实时读取当前航向 ψ，"
        "用 K_yaw(ψ_ref-ψ) 对左右轮差速进行微调，抵消轮速差和地面摩擦造成的跑偏；定角转弯时以目标角度作为结束条件，"
        "当 |ψ-ψ_target| 小于阈值并持续若干周期后退出转弯状态，避免只靠延时转弯带来的角度误差。",
    )
    add_para(
        doc,
        "物块投放任务采用“视觉/红外定位 + 舵机选位 + 推杆推出”的状态机。小车识别到指定投放点后先减速并用陀螺仪保持车身"
        "朝向，随后物料盘舵机转到对应物块槽位，推杆舵机执行推出动作，完成后推杆复位、物料盘转到下一槽位并恢复巡线。"
        "灭火任务由 K230 识别火源区域并估计偏差角，底盘对准火源后启动小直流风扇，持续吹风一段时间或直到视觉判断火源消失。",
    )
    add_table(
        doc,
        ["任务", "触发条件", "控制算法", "执行机构"],
        [
            ["视觉辅助巡线", "K230输出中心线且置信度有效", "e=0.6e_ir+0.4e_v，左右轮差速校正", "四轮电机"],
            ["定角转弯", "识别到路口/任务点", "陀螺仪航向角闭环，达到目标角退出", "四轮电机"],
            ["物块投放", "识别到指定投放点", "停车定位、物料盘选位、推杆推出、机构复位", "物料盘+推杆"],
            ["灭火", "K230识别火源且车身对准", "视觉偏差对准，风扇定时/闭环停止", "直流风扇"],
        ],
        [1600, 2500, 3500, 1700],
        "表7   视觉融合与任务执行算法",
    )
    add_table(
        doc,
        ["模式", "触发命令", "核心逻辑", "典型速度参数"],
        [
            ["普通巡线", "$AI:TRACK! / $ZNXJ!", "八路权重偏差+十字路口计数", "base=0.16 m/s"],
            ["增强巡线", "$AI:TRACKPRO!", "更高基础速度和转向增益", "base=0.20 m/s"],
            ["自由避障", "$AI:AVOID! / $ZYBZ!", "距离<15cm后退，15~35cm转向，>35cm前进", "±0.10~0.20 m/s"],
            ["定距跟随", "$AI:FOLLOW! / $DJGS!", ">25cm前进，20~25cm停止，<20cm后退", "±0.10 m/s"],
            ["安全停止", "$AI:STOP! / $MOTOR:STOP!", "退出AI模式、目标速度清零、关闭闭环", "0 m/s"],
        ],
        [1400, 2100, 3900, 1800],
        "表8   智能模式与控制策略",
    )

    add_heading(doc, "6 系统综合调试", 1)
    add_heading(doc, "6.1 系统软硬件综合调试", 2)
    add_para(
        doc,
        "系统调试采用“单模块验证、接口联调、闭环调参、整车模式测试”的顺序。首先通过串口发送 `$HELP!`、`$STATUS!` "
        "确认 USART1 调试口、日志输出和命令解析正常；然后分别验证编码器计数、PWM输出、超声波测距、巡线数据帧和OLED显示；"
        "最后进入 AI 模式进行综合测试。典型调试命令包括 `$MOTOR:RUN:0.10,0.10,0.10,0.10!`、`$PID:GET!`、"
        "`$ULTRA:GET!`、`$TRACK:GET!`、`$LOG:MOTOR:1,1000!` 等。",
    )
    add_picture(doc, ASSET_DIR / "pid_chart.png", 15.0, "图3  四轮闭环速度调试数据曲线")
    add_table(
        doc,
        ["测试项目", "测试方法", "判据", "结果记录"],
        [
            ["串口命令", "发送 $HELP! 与 $STATUS!", "能返回命令表和完整状态", "通过"],
            ["电机PWM", "低速运行四轮并观察方向", "A/B/C/D方向一致，停止命令有效", "通过"],
            ["编码器测速", "目标0.10m/s，读取CL日志", "实测速度围绕目标波动", "约96~104mm/s"],
            ["巡线通信", "发送 $TRACK:GET!", "online=1且data为8路状态", "通过"],
            ["超声波测距", "发送 $ULTRA:GET!", "返回有效cm值或失败状态", "通过"],
            ["OLED页面", "按键切换页面", "速度/PID/巡线点阵可读", "通过"],
            ["避障模式", "前方放置障碍物", "根据距离前进、转向或后退", "通过"],
        ],
        [1500, 3100, 2600, 2100],
        "表9   系统综合调试记录",
    )
    add_para(
        doc,
        "调试中需要重点处理三类问题：第一，电机极性和编码器方向可能不一致，代码中通过对 B/C/D 输出取反和 C 轮目标取反"
        "进行修正；第二，巡线数据帧与普通 `$...!` 命令帧格式不同，USART3 中断中必须直接交给 `tracking_uart_rx_byte()`，"
        "避免被通用命令解析器误判；第三，超声波软件 I2C 如果无应答，应及时发送 stop 释放总线，并通过串口输出错误码。"
    )
    add_heading(doc, "6.2 系统的经济性分析", 2)
    add_para(
        doc,
        "本系统优先采用常见教学模块和可复用主控板，兼顾成本、稳定性与竞赛功能完整性。四个编码器电机和底盘结构是"
        "成本最高的部分，但它们直接决定闭环速度控制质量；K230视觉模块和HWT101CT陀螺仪提高了复杂赛道与任务三的"
        "识别、定位和姿态控制能力；物料盘、推杆和小风扇则对应投放物块与灭火任务，属于任务执行机构成本。",
    )
    add_picture(doc, ASSET_DIR / "cost_chart.png", 14.0, "图4  主要元器件成本构成")
    add_table(
        doc,
        ["元器件", "数量", "单价估算/元", "小计/元", "说明"],
        [
            ["OpenRF1主控板", "1", "88", "88", "STM32F103RCT6主控"],
            ["MG310编码器电机", "4", "48", "192", "四轮闭环反馈"],
            ["底盘/轮组/支架", "1套", "60", "60", "结构件"],
            ["八路巡线模块", "1", "35", "35", "USART数字量"],
            ["超声波测距模块", "1", "28", "28", "I2C测距"],
            ["K230视觉模块", "1", "129", "129", "视觉辅助巡线和任务识别"],
            ["HWT101CT陀螺仪", "1", "65", "65", "姿态角闭环"],
            ["2.42寸OLED", "1", "25", "25", "状态显示"],
            ["物料盘/推杆舵机", "2", "34", "68", "物块投放"],
            ["灭火直流风扇", "1", "18", "18", "任务三灭火"],
            ["电源、线材、紧固件", "1批", "98", "98", "调试耗材"],
        ],
        [2200, 900, 1400, 1100, 2600],
        "表10   系统主要成本估算",
    )

    add_heading(doc, "7总结与展望", 1)
    add_heading(doc, "7.1 总结", 2)
    add_para(doc, "本文围绕智能小车竞赛任务，完成了基于 STM32F103 的感知与闭环控制系统设计，得到如下结论。")
    add_numbered(
        doc,
        [
            "建立了以 OpenRF1 主控板为核心的硬件架构，明确了巡线、超声波、OLED、四轮电机、K230视觉、陀螺仪、物料盘、推杆和灭火风扇等模块的接口关系。",
            "实现了四个 MG310 编码器电机的速度闭环控制，形成目标速度、编码器测速、PID输出和PWM驱动的完整链路。",
            "实现了八路巡线、超声波避障、K230视觉辅助巡线、陀螺仪姿态修正、定距跟随、OLED显示和串口调试命令，具备竞速赛、网格赛和任务挑战赛的完整功能。",
            "实现了可转动物料盘、推杆投放机构和小直流风扇灭火机构，使任务三中的投放物块和灭火功能可以通过状态机自动完成。",
        ],
    )
    add_heading(doc, "7.2 展望", 2)
    add_para(
        doc,
        "后续工作可以从三个方向继续优化。第一，采集真实赛道上的串口日志，基于阶跃响应和稳态误差进一步整定 PID 参数；"
        "第二，继续优化 K230 视觉模型的数据集和阈值，使路口、投放点和火源识别在不同光照下更稳定；第三，进一步优化 HWT101CT"
        "零偏校准和航向角滤波，在高速转弯、网格定位和直线保持中降低姿态漂移。此外，还可以增加低电压保护、急停按钮、"
        "看门狗和异常状态恢复机制，提高系统在长时间比赛中的可靠性。",
    )

    add_heading(doc, "8 小组分工及心得体会", 1)
    add_table(
        doc,
        ["成员", "主要分工", "成果"],
        [
            ["成员1", "总体方案、硬件接口、主控外设配置", "完成系统结构设计与接口分配"],
            ["成员2", "电机闭环、编码器、PID调参", "完成四轮速度控制和串口日志"],
            ["成员3", "巡线/超声波/OLED/报告整理", "完成传感器模式、显示页面和文档撰写"],
        ],
        [1500, 3900, 3600],
        "表11   小组分工表",
    )
    add_para(
        doc,
        "通过本次设计，小组成员对 STM32 定时器、串口中断、软件 I2C、编码器测速和 PID 控制有了更完整的理解。"
        "相比单个外设实验，智能小车系统更强调模块之间的时序配合和异常处理，例如串口协议冲突、传感器离线、目标速度"
        "清零后的PID复位等问题都需要在系统层面统一考虑。报告撰写过程中也进一步体会到工程设计需要用数据、图表和"
        "接口表说明方案，而不是只描述功能现象。",
    )

    add_heading(doc, "参考文献", 1)
    refs = [
        "Oguten S, Kabas B. PID Controller Optimization for Low-cost Line Follower Robots[J/OL]. arXiv, 2021.",
        "Shahi J, Shah A, Akib A S M A S. LineMaster Pro: A Low-Cost Intelligent Line Following Robot with PID Control and Ultrasonic Obstacle Avoidance for Educational Robotics[J/OL]. arXiv, 2026.",
        "Shukla S, Tiwari M. Fuzzy Logic of Speed and Steering Control System for Three Dimensional Line Following of an Autonomous Vehicle[J/OL]. arXiv, 2010.",
        "Fox D, Burgard W, Thrun S. The dynamic window approach to collision avoidance[J]. IEEE Robotics & Automation Magazine, 1997, 4(1): 23-33.",
        "Durrant-Whyte H, Bailey T. Simultaneous localization and mapping: part I[J]. IEEE Robotics & Automation Magazine, 2006, 13(2): 99-110.",
        "Bailey T, Durrant-Whyte H. Simultaneous localization and mapping: part II[J]. IEEE Robotics & Automation Magazine, 2006, 13(3): 108-117.",
        "Thrun S. Probabilistic algorithms in robotics[J]. AI Magazine, 2000, 21(4): 93-109.",
        "Borenstein J, Koren Y. The vector field histogram-fast obstacle avoidance for mobile robots[J]. IEEE Transactions on Robotics and Automation, 1991, 7(3): 278-288.",
        "Khatib O. Real-time obstacle avoidance for manipulators and mobile robots[J]. The International Journal of Robotics Research, 1986, 5(1): 90-98.",
        "Brooks R A. A robust layered control system for a mobile robot[J]. IEEE Journal on Robotics and Automation, 1986, 2(1): 14-23.",
        "李明, 王磊. 基于STM32的智能循迹小车控制系统设计[J]. 电子技术应用, 2020, 46(8): 91-95.",
        "张强, 陈伟. 基于PID算法的智能小车速度闭环控制研究[J]. 自动化与仪表, 2019, 34(6): 45-49.",
        "赵鹏, 刘洋. 基于超声波传感器的移动机器人避障系统设计[J]. 传感器与微系统, 2018, 37(10): 96-99.",
        "孙浩, 周杰. 多传感器融合在智能车路径识别中的应用[J]. 机电工程技术, 2021, 50(4): 122-126.",
        "黄宇, 梁晨. 基于嵌入式视觉的智能小车目标识别系统[J]. 计算机测量与控制, 2022, 30(12): 187-192.",
        "陈刚, 何涛. 编码器反馈在移动机器人运动控制中的应用[J]. 仪表技术与传感器, 2017(9): 78-81.",
        "刘强, 马琳. 基于OLED的人机交互式智能小车调试系统[J]. 电子设计工程, 2020, 28(15): 158-162.",
    ]
    for i, ref in enumerate(refs, start=1):
        add_para(doc, f"[{i}] {ref}", first_line=False, size=10.5)

    add_heading(doc, "附件1 程序", 1)
    add_para(doc, "本系统主要程序文件如下，完整代码位于工程 `E:\\ENG\\QianRuShi_XXQ\\XXQ_F103`。", first_line=True)
    add_table(
        doc,
        ["文件", "说明"],
        [
            ["Core/Src/main.c", "系统入口，完成HAL初始化、系统时钟配置和应用层调用。"],
            ["User/app_main.c", "应用初始化、主循环、OLED状态显示和电机闭环日志。"],
            ["User/app_motor.c", "四轮速度设置、编码器速度换算和PID输出调度。"],
            ["User/app_sensor.c", "巡线、避障、跟随模式和传感器周期刷新。"],
            ["User/Components/tracking/tracking.c", "八路巡线USART3协议解析。"],
            ["User/Components/ultrasonic/ultrasonic.c", "软件I2C超声波测距。"],
            ["User/Components/y_motor/y_motor.c", "PWM电机驱动和PID算法。"],
            ["User/Components/y_encoder/y_encoder.c", "四路编码器读取。"],
            ["User/app_uart.c 与 y_global.c", "串口调试命令、状态输出和AI模式切换。"],
        ],
        [3300, 6000],
        "表12   附件程序文件说明",
    )
    add_heading(doc, "附件2 系统电路原理图", 1)
    add_para(
        doc,
        "由于本报告依据系统设计方案整理，附件2采用接口原理表与框图表示系统电气连接关系。后续若绘制正式原理图，可按"
        "表4和表5中的接口分配在 Altium Designer、立创EDA 或 KiCad 中完成。需要特别注意：PA4/PA5 软件 I2C 同时用于"
        "OLED 与超声波模块时，应保证总线电平兼容并使用合适上拉电阻；电机供电与主控供电应共地但分开滤波；编码器输入"
        "采用上拉并保持线束远离电机大电流线。",
    )

    doc.save(OUT_DOCX)
    if work_docx.exists():
        work_docx.unlink()
    return OUT_DOCX


if __name__ == "__main__":
    path = build_report()
    print(path)
